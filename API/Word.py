import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


import os


class Word:
    def __init__(self, filename, path):
        self.filename = filename
        self.path = path
        self.file = self.path + f"\\{self.filename}.docx"
        self.document = self.newDocument()
        self.createStyles()

    def newDocument(self):
        document = docx.Document()
        return document

    def saveDocument(self):
        self.document = self.document.save(self.file)
        print("Saved.")

    def openDocument(self):
        try:
            document = docx.Document(self.file)
            print(f"Successfully Opened {self.file}")
            return document
        except:
            raise f"Could not open file {self.file}."

    def isExist(self):
        if os.path.exists(self.file):
            return True
        else:
            return False

    def setTopic(self, topic):
        self.document.add_heading(f"{topic}", 0)
        print("Created topic")

    def addParagraph(self):
        p = self.document.add_paragraph()
        return p

    def addHeader(self):
        return self.document.sections[0].header

    def createStyles(self):
        obj_styles = self.document.styles
        obj_charstyle = obj_styles.add_style("small", WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(12)
        obj_font.name = "Calibri (Body)"

        obj_charstyle = obj_styles.add_style("big", WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(14)
        obj_font.name = "Calibri (Body)"

        obj_charstyle = obj_styles.add_style("large", WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(22)
        obj_font.name = "Calibri (Body)"

    def addText(
        self,
        text,
        style="small",
        header=False,
        align=0,
        bold=False,
        bullet=False,
        space_before=0,
        space_after=0,
        tabs=0,
    ):
        if header:
            p = self.document.sections[0].header.paragraphs[0]
        else:
            p = self.document.add_paragraph()
            # p.paragraph_format.line_spacing = Pt(11)
        # p.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if bullet:
            p.paragraph_format.left_indent = Inches(0.75)
            p.align = 0
            p.style = "List Bullet"
        else:
            p.alignment = align
        p.add_run(text, style=style).bold = bold

    def addBreak(self, pos):
        if self.document.paragraphs:
            para = self.document.paragraphs[-1]
        else:
            para = self.document.add_paragraph()

        # para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # para.paragraph_format.line_spacing = Pt(0)

        p = para._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        pPr.insert_element_before(
            pBdr,
            "w:shd",
            "w:tabs",
            "w:suppressAutoHyphens",
            "w:kinsoku",
            "w:wordWrap",
            "w:overflowPunct",
            "w:topLinePunct",
            "w:autoSpaceDE",
            "w:autoSpaceDN",
            "w:bidi",
            "w:adjustRightInd",
            "w:snapToGrid",
            "w:spacing",
            "w:ind",
            "w:contextualSpacing",
            "w:mirrorIndents",
            "w:suppressOverlap",
            "w:jc",
            "w:textDirection",
            "w:textAlignment",
            "w:textboxTightWrap",
            "w:outlineLvl",
            "w:divId",
            "w:cnfStyle",
            "w:rPr",
            "w:sectPr",
            "w:pPrChange",
        )
        bottom = OxmlElement(f"w:{pos}")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        # para.add_run(""),

    def addHeaderText(self, header, text, font="Calibri", size=22, align=0, bold=True):
        header = header.paragraphs[0]
        header.style.font.name = font
        header.style.font.size = Pt(size)
        header.style.font.color.rgb = RGBColor(0, 0, 0)
        header.alignment = align
        # header.style.bold = bold
        header.add_run(text + "\n").bold = bold

    def replaceInParagraph(self, old, new):
        for paragraph in self.document.paragraphs:
            if old in paragraph.text.upper():
                print("Found a match between " + old + " and " + new)
                print("Old Text: " + paragraph.text)
                paragraph.text = paragraph.text.replace(old, new)
                print("New Text: " + paragraph.text)

                return
        for table in self.document.tables:
            # print(old)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old in para.text.upper():
                            print(f"Found a match HERE between {old} and {new}")
                            print("Old Text: " + para.text)
                            print("Old: ", old)
                            para.text = para.text.upper().replace(
                                old.upper(), new.upper()
                            )
                            print("New Text: " + para.text)
                            return
        print("Couldnt find a match " + old)
