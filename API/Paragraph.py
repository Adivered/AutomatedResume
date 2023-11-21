from API.Text import Text
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


class Paragraph:
    def __init__(self, document, bullet="false"):
        self._document = document
        self._paragraph = self._document.add_paragraph()
        self._bullet = bullet

        if bullet:
            self._paragraph.style = "List Bullet"

    @property
    def align(self):
        return self._align

    @align.setter
    def align(self, value):
        self._align = value.lower()
        if self._align == "left":
            self._paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif self._align == "right":
            self._paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif self._align == "center":
            self._paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Handle other alignment values as needed

    @property
    def bold(self):
        return self._bold

    @bold.setter
    def bold(self, value):
        self._bold = value

    @property
    def space_before(self):
        return self._space_before

    @space_before.setter
    def space_before(self, value):
        self._space_before = Pt(value)
        self._paragraph.paragraph_format.space_before = self._space_before

    @property
    def space_after(self):
        return self._space_after

    @space_after.setter
    def space_after(self, value):
        self._space_after = Pt(value)
        self._paragraph.paragraph_format.space_after = self._space_after

    @property
    def style(self):
        return self._style

    @style.setter
    def style(self, value):
        self._style = value

    def add_text(self, text, style=None):
        text_obj = Text(text, style)
        if self._bullet:
            self._paragraph.paragraph_format.left_indent = Inches(0.75)
            # self._paragraph.add_run(text_obj.text, style=self._style)

        if style:
            self._paragraph.add_run(text_obj.text, style=style)
        else:
            self._paragraph.add_run(text_obj.text)

    def add_tab(self):
        run = self._paragraph.add_run()
        run.add_tab()

    def add_break_line(self, pos):
        p = self._paragraph
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        # p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        pPr.insert_element_before(
            pBdr,
            "w:shd",
            "w:tabs",
            "w:suppressAutoHyphens",
            "w:kinsoku",
            "w:wordWrap",
            "w:overflowPunct",
            "w:topLinePunct",
            "w:autoSpaceDE",
            "w:autoSpaceDN",
            "w:bidi",
            "w:adjustRightInd",
            "w:snapToGrid",
            "w:spacing",
            "w:ind",
            "w:contextualSpacing",
            "w:mirrorIndents",
            "w:suppressOverlap",
            "w:jc",
            "w:textDirection",
            "w:textAlignment",
            "w:textboxTightWrap",
            "w:outlineLvl",
            "w:divId",
            "w:cnfStyle",
            "w:rPr",
            "w:sectPr",
            "w:pPrChange",
        )
        bottom = OxmlElement(f"w:{pos}")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
