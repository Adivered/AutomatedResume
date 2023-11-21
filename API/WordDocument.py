from docx import Document as docxDocument
from os import path as ospath
from API.Paragraph import Paragraph


class WordDocument:
    def __init__(self, filename, path):
        self._filename = filename
        self._path = path
        self._file = self._path + f"\\{self._filename}.docx"
        self._document = self.new_document()

    def new_document(self):
        document = docxDocument()
        return document

    def save_document(self):
        self._document.save(self._file)
        print("Saved.")

    def open_document(self):
        try:
            document = docxDocument(self.file)
            print(f"Successfully Opened {self.file}")
            return document
        except:
            raise f"Could not open file {self.file}."

    def is_exist(self):
        return ospath.exists(self.file)

    def add_paragraph(self, bullet=False):
        return Paragraph(self._document, bullet)

    def add_style(self, style):
        """
        Add a custom style to the document.

        :param style: A Style object defining the custom style.
        """
        obj_styles = self._document.styles
        obj_charstyle = obj_styles.add_style(style._name, style._character)
        obj_font = obj_charstyle.font
        obj_font.size = style._font_size
        obj_font.name = style._font_name
        obj_font.bold = style._font_bold
        obj_font.underline = style._font_underline

        # Apply other style attributes from the Style object
        if style._font_color:
            obj_font.color.rgb = style._font_color
